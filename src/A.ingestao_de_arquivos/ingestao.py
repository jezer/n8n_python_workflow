import os
import time
from pathlib import Path
from typing import List, Dict, Any, Optional, Callable, Union
from dataclasses import dataclass, asdict
from enum import Enum
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
import mimetypes
from contextlib import contextmanager

# Importações condicionais com tratamento de erro
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import docx
    import mammoth
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from markdownify import markdownify as md
    import markdown
    HAS_MARKDOWN = True
except ImportError:
    HAS_MARKDOWN = False

import json
import pprint


class SupportedFormat(Enum):
    """Enumeration dos formatos suportados"""
    TXT = ".txt"
    MD = ".md" 
    DOCX = ".docx"
    XLS = ".xls"
    XLSX = ".xlsx"
    PDF = ".pdf"
    JSON = ".json"


class ReturnFormat(Enum):
    """Enumeration dos formatos de retorno"""
    MARKDOWN = "markdown"
    TEXT = "text"
    HTML = "html"


@dataclass
class DocumentMetadata:
    """Estrutura de dados para metadados do documento"""
    file_size_bytes: int
    file_format: str
    file_path: str
    processing_time_seconds: float
    content_size_chars: int
    content_lines: int
    file_hash: str
    mime_type: str
    last_modified: float
    encoding_detected: Optional[str] = None
    ocr_applied: bool = False
    llm_processed: bool = False
    error_count: int = 0


@dataclass
class ProcessedDocument:
    """Estrutura de dados para documento processado"""
    filename: str
    format: str
    content: str
    metadata: DocumentMetadata
    
    def to_dict(self) -> Dict[str, Any]:
        """Converte para dicionário mantendo compatibilidade com API original"""
        return {
            "nome_arquivo": self.filename,
            "formato": self.format,
            "conteudo": self.content,
            "metadados": asdict(self.metadata)
        }


class ProcessingError(Exception):
    """Exceção customizada para erros de processamento"""
    pass


class UnsupportedFormatError(ProcessingError):
    """Exceção para formatos não suportados"""
    pass


class FileExtractionError(ProcessingError):
    """Exceção para erros de extração de arquivo"""
    pass


class IngestaoDeArquivos:
    """
    Classe aprimorada para ingestão de arquivos com melhor arquitetura,
    tratamento de erros, performance e extensibilidade.
    """

    def __init__(
        self,
        tesseract_path: Optional[str] = None,
        output_path: Optional[str] = None,
        log_level: int = logging.INFO,
        language: str = "pt",
        max_workers: int = 4,
        ocr_timeout: int = 300,
        max_file_size_mb: int = 100,
        chunk_size: int = 8192,
        enable_file_validation: bool = True
    ):
        """
        Inicializa a classe de ingestão com configurações aprimoradas.
        
        Args:
            tesseract_path: Caminho para executável do Tesseract
            output_path: Diretório de saída para arquivos processados
            log_level: Nível de logging
            language: Idioma para OCR
            max_workers: Número máximo de threads paralelas
            ocr_timeout: Timeout para operações OCR em segundos
            max_file_size_mb: Tamanho máximo de arquivo em MB
            chunk_size: Tamanho do chunk para leitura de arquivos
            enable_file_validation: Habilita validação de arquivos
        """
        self._setup_dependencies(tesseract_path)
        self._setup_configuration(
            output_path, language, max_workers, ocr_timeout, 
            max_file_size_mb, chunk_size, enable_file_validation
        )
        self._setup_logging(log_level)
        self._initialize_extractors()

    def _setup_dependencies(self, tesseract_path: Optional[str]) -> None:
        """Configura dependências externas"""
        if tesseract_path and HAS_OCR:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Verifica dependências críticas
        self._check_dependencies()

    def _setup_configuration(
        self, output_path: Optional[str], language: str, max_workers: int,
        ocr_timeout: int, max_file_size_mb: int, chunk_size: int,
        enable_file_validation: bool
    ) -> None:
        """Configura parâmetros da classe"""
        self.supported_formats = [fmt.value for fmt in SupportedFormat]
        self.output_path = Path(output_path) if output_path else None
        if self.output_path:
            self.output_path.mkdir(parents=True, exist_ok=True)
        
        self.language = language
        self.max_workers = max_workers
        self.ocr_timeout = ocr_timeout
        self.max_file_size_bytes = max_file_size_mb * 1024 * 1024
        self.chunk_size = chunk_size
        self.enable_file_validation = enable_file_validation

    def _setup_logging(self, log_level: int) -> None:
        """Configura sistema de logging"""
        logging.basicConfig(
            level=log_level,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.StreamHandler(),
                logging.FileHandler('ingestao.log', encoding='utf-8')
            ]
        )
        self.logger = logging.getLogger(self.__class__.__name__)

    def _check_dependencies(self) -> None:
        """Verifica dependências opcionais e registra avisos"""
        dependencies = {
            "PyMuPDF (fitz)": HAS_PYMUPDF,
            "OCR (pytesseract, PIL, pdf2image)": HAS_OCR,
            "DOCX (python-docx, mammoth)": HAS_DOCX,
            "Excel (pandas)": HAS_PANDAS,
            "Markdown (markdownify, markdown)": HAS_MARKDOWN
        }
        
        missing = [name for name, available in dependencies.items() if not available]
        if missing:
            self.logger.warning(f"Dependências opcionais não disponíveis: {', '.join(missing)}")

    def _initialize_extractors(self) -> None:
        """Inicializa mapeamento de extratores por formato"""
        self._extractors = {
            SupportedFormat.TXT.value: self._extract_txt,
            SupportedFormat.MD.value: self._extract_txt,
            SupportedFormat.DOCX.value: self._extract_docx,
            SupportedFormat.XLS.value: self._extract_excel,
            SupportedFormat.XLSX.value: self._extract_excel,
            SupportedFormat.PDF.value: self._extract_pdf,
            SupportedFormat.JSON.value: self._extract_json
        }

    def _validate_file(self, file_path: Path) -> None:
        """
        Valida arquivo antes do processamento.
        
        Raises:
            ProcessingError: Se arquivo não passar na validação
        """
        if not file_path.exists():
            raise ProcessingError(f"Arquivo não encontrado: {file_path}")
        
        if not file_path.is_file():
            raise ProcessingError(f"Caminho não é um arquivo: {file_path}")
        
        if file_path.stat().st_size > self.max_file_size_bytes:
            raise ProcessingError(
                f"Arquivo muito grande: {file_path.stat().st_size} bytes "
                f"(máximo: {self.max_file_size_bytes} bytes)"
            )
        
        ext = file_path.suffix.lower()
        if ext not in self.supported_formats:
            raise UnsupportedFormatError(f"Formato não suportado: {ext}")

    def _detect_encoding(self, file_path: Path) -> str:
        """
        Detecta encoding do arquivo texto.
        
        Returns:
            String do encoding detectado
        """
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(min(self.chunk_size, 10000))
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except ImportError:
            # Fallback para detecção manual
            encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        f.read(1000)  # Testa leitura
                    return encoding
                except UnicodeDecodeError:
                    continue
            return 'utf-8'

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calcula hash SHA256 do arquivo"""
        hash_sha256 = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(self.chunk_size), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            self.logger.warning(f"Erro ao calcular hash de {file_path}: {e}")
            return "unknown"

    @contextmanager
    def _performance_timer(self):
        """Context manager para medir tempo de processamento"""
        start_time = time.time()
        yield
        end_time = time.time()
        self.processing_time = end_time - start_time

    def process_file(
        self,
        file_path: Union[str, Path],
        save_markdown: bool = False,
        return_format: Union[str, ReturnFormat] = ReturnFormat.MARKDOWN,
        ocr_func: Optional[Callable] = None,
        llm_func: Optional[Callable] = None
    ) -> Optional[ProcessedDocument]:
        """
        Processa um único arquivo com melhor tratamento de erro e performance.
        
        Args:
            file_path: Caminho do arquivo
            save_markdown: Se deve salvar em formato markdown
            return_format: Formato de retorno do conteúdo
            ocr_func: Função personalizada para OCR
            llm_func: Função personalizada para pós-processamento LLM
            
        Returns:
            ProcessedDocument ou None se erro
        """
        file_path = Path(file_path)
        
        # Configura funções padrão
        ocr_func = ocr_func or (pytesseract.image_to_string if HAS_OCR else None)
        llm_func = llm_func or self._default_llm_processing
        
        # Converte return_format para enum se necessário
        if isinstance(return_format, str):
            try:
                return_format = ReturnFormat(return_format)
            except ValueError:
                self.logger.warning(f"Formato de retorno inválido: {return_format}, usando markdown")
                return_format = ReturnFormat.MARKDOWN

        try:
            with self._performance_timer():
                # Validação
                if self.enable_file_validation:
                    self._validate_file(file_path)
                
                # Coleta metadados básicos
                stat = file_path.stat()
                file_hash = self._calculate_file_hash(file_path)
                mime_type = mimetypes.guess_type(str(file_path))[0] or "unknown"
                
                # Extração de conteúdo
                ext = file_path.suffix.lower()
                extractor = self._extractors.get(ext)
                
                if not extractor:
                    raise UnsupportedFormatError(f"Extrator não encontrado para: {ext}")
                
                # Extrai conteúdo
                extraction_result = extractor(file_path, ocr_func)
                
                if isinstance(extraction_result, tuple):
                    content, ocr_applied = extraction_result
                else:
                    content, ocr_applied = extraction_result, False
                
                # Pós-processamento LLM
                processed_content, llm_processed = self._apply_llm_processing(content, llm_func)
                
                # Converte para formato solicitado
                final_content = self._convert_content_format(processed_content, return_format)
                
                # Cria metadados
                metadata = DocumentMetadata(
                    file_size_bytes=stat.st_size,
                    file_format=ext[1:],
                    file_path=str(file_path),
                    processing_time_seconds=self.processing_time,
                    content_size_chars=len(final_content),
                    content_lines=final_content.count('\n') + 1,
                    file_hash=file_hash,
                    mime_type=mime_type,
                    last_modified=stat.st_mtime,
                    ocr_applied=ocr_applied,
                    llm_processed=llm_processed
                )
                
                # Salva markdown se solicitado
                if save_markdown and self.output_path:
                    self._save_markdown(file_path, processed_content)
                
                # Cria documento processado
                document = ProcessedDocument(
                    filename=file_path.name,
                    format=ext[1:],
                    content=final_content,
                    metadata=metadata
                )
                
                self.logger.info(
                    f"Arquivo processado com sucesso: {file_path.name} "
                    f"({stat.st_size} bytes em {self.processing_time:.2f}s)"
                )
                
                return document
                
        except Exception as e:
            self.logger.error(f"Erro ao processar {file_path}: {type(e).__name__}: {e}")
            return None

    def process_directory(
        self,
        dir_path: Union[str, Path],
        save_markdown: bool = False,
        filter_ext: Optional[List[str]] = None,
        min_size: int = 0,
        max_size: Optional[int] = None,
        return_format: Union[str, ReturnFormat] = ReturnFormat.MARKDOWN,
        recursive: bool = False,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> List[ProcessedDocument]:
        """
        Processa arquivos em diretório com opções avançadas.
        
        Args:
            dir_path: Caminho do diretório
            save_markdown: Se deve salvar em markdown
            filter_ext: Lista de extensões para filtrar
            min_size: Tamanho mínimo em bytes
            max_size: Tamanho máximo em bytes
            return_format: Formato de retorno
            recursive: Se deve processar recursivamente
            progress_callback: Callback para progresso (atual, total)
            
        Returns:
            Lista de documentos processados
        """
        dir_path = Path(dir_path)
        
        if not dir_path.exists() or not dir_path.is_dir():
            self.logger.warning(f"Diretório não encontrado: {dir_path}")
            return []

        # Coleta arquivos
        files = self._collect_files(dir_path, filter_ext, min_size, max_size, recursive)
        
        if not files:
            self.logger.info(f"Nenhum arquivo encontrado em: {dir_path}")
            return []

        self.logger.info(f"Processando {len(files)} arquivos com {self.max_workers} workers")
        
        documents = []
        processed_count = 0
        
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submete tarefas
            future_to_file = {
                executor.submit(
                    self.process_file, 
                    file_path, 
                    save_markdown, 
                    return_format
                ): file_path
                for file_path in files
            }
            
            # Coleta resultados
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    document = future.result()
                    if document:
                        documents.append(document)
                    
                    processed_count += 1
                    
                    # Callback de progresso
                    if progress_callback:
                        progress_callback(processed_count, len(files))
                        
                except Exception as e:
                    self.logger.error(f"Erro no processamento paralelo de {file_path}: {e}")
        
        self.logger.info(f"Processamento concluído: {len(documents)} documentos de {len(files)} arquivos")
        return documents

    def _collect_files(
        self,
        dir_path: Path,
        filter_ext: Optional[List[str]],
        min_size: int,
        max_size: Optional[int],
        recursive: bool
    ) -> List[Path]:
        """Coleta arquivos aplicando filtros"""
        files = []
        
        # Padrão de busca
        pattern = "**/*" if recursive else "*"
        
        for file_path in dir_path.glob(pattern):
            if not file_path.is_file():
                continue
            
            # Filtro por extensão
            ext = file_path.suffix.lower()
            if filter_ext:
                if ext not in filter_ext:
                    continue
            elif ext not in self.supported_formats:
                continue
            
            # Filtro por tamanho
            file_size = file_path.stat().st_size
            if file_size < min_size:
                continue
            if max_size and file_size > max_size:
                continue
            
            files.append(file_path)
        
        return files

    def _apply_llm_processing(
        self, 
        content: str, 
        llm_func: Optional[Callable]
    ) -> tuple[str, bool]:
        """Aplica processamento LLM se função fornecida"""
        if llm_func and content.strip():
            try:
                processed = llm_func(content)
                return processed, True
            except Exception as e:
                self.logger.warning(f"Erro no processamento LLM: {e}")
                return content, False
        return content, False

    def _convert_content_format(self, content: str, format_type: ReturnFormat) -> str:
        """Converte conteúdo para formato solicitado"""
        if format_type == ReturnFormat.MARKDOWN:
            return content
        elif format_type == ReturnFormat.TEXT:
            return self._markdown_to_text(content)
        elif format_type == ReturnFormat.HTML:
            return self._markdown_to_html(content)
        else:
            return content

    def _save_markdown(self, original_path: Path, content: str) -> None:
        """Salva conteúdo em arquivo markdown"""
        try:
            output_file = self.output_path / (original_path.stem + ".md")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(content)
            self.logger.debug(f"Markdown salvo: {output_file}")
        except Exception as e:
            self.logger.error(f"Erro ao salvar markdown: {e}")

    # ===== MÉTODOS DE EXTRAÇÃO =====

    def _extract_txt(self, file_path: Path, ocr_func: Optional[Callable] = None) -> str:
        """Extrai texto de arquivo TXT/MD com detecção de encoding"""
        encoding = self._detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Converte para markdown se for arquivo MD
            if file_path.suffix.lower() == '.md':
                return content
            else:
                return md(content) if HAS_MARKDOWN else content
                
        except Exception as e:
            raise FileExtractionError(f"Erro ao ler arquivo texto: {e}")

    def _extract_docx(self, file_path: Path, ocr_func: Optional[Callable] = None) -> str:
        """Extrai texto de arquivo DOCX"""
        if not HAS_DOCX:
            raise ProcessingError("Dependências DOCX não disponíveis")
        
        try:
            # Tenta usar mammoth primeiro (melhor formatação)
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                if result.value.strip():
                    return result.value
        except Exception as e:
            self.logger.warning(f"Falha com mammoth, tentando python-docx: {e}")
        
        try:
            # Fallback para python-docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            return md(content) if HAS_MARKDOWN else content
        except Exception as e:
            raise FileExtractionError(f"Erro ao processar DOCX: {e}")

    def _extract_excel(self, file_path: Path, ocr_func: Optional[Callable] = None) -> str:
        """Extrai dados de arquivo Excel"""
        if not HAS_PANDAS:
            raise ProcessingError("Dependência pandas não disponível")
        
        try:
            # Lê todas as abas
            dfs = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            markdown_parts = []
            for sheet_name, df in dfs.items():
                # Remove colunas/linhas completamente vazias
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                if not df.empty:
                    markdown_parts.append(f"## Planilha: {sheet_name}\n\n{df.to_markdown(index=False)}")
            
            return "\n\n".join(markdown_parts)
            
        except Exception as e:
            raise FileExtractionError(f"Erro ao processar Excel: {e}")

    def _extract_pdf(self, file_path: Path, ocr_func: Optional[Callable] = None) -> tuple[str, bool]:
        """Extrai texto de PDF com OCR quando necessário"""
        if not HAS_PYMUPDF:
            raise ProcessingError("Dependência PyMuPDF não disponível")
        
        try:
            # Tenta extrair texto nativo primeiro
            doc = fitz.open(str(file_path))
            text_content = ""
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                text_content += page_text
            
            doc.close()
            
            # Se tem texto suficiente, usa extração nativa
            if len(text_content.strip()) > 50:  # Threshold mínimo
                return md(text_content) if HAS_MARKDOWN else text_content, False
            
            # Caso contrário, usa OCR
            return self._extract_pdf_ocr(file_path, ocr_func)
            
        except Exception as e:
            self.logger.warning(f"Erro na extração nativa de PDF, tentando OCR: {e}")
            return self._extract_pdf_ocr(file_path, ocr_func)

    def _extract_pdf_ocr(self, file_path: Path, ocr_func: Optional[Callable] = None) -> tuple[str, bool]:
        """Extrai texto de PDF usando OCR"""
        if not HAS_OCR:
            raise ProcessingError("Dependências OCR não disponíveis")
        
        if not ocr_func:
            raise ProcessingError("Função OCR não fornecida")
        
        try:
            # Converte PDF para imagens
            images = convert_from_path(str(file_path))
            
            text_parts = []
            for i, image in enumerate(images):
                try:
                    # Aplica OCR com timeout
                    page_text = ocr_func(image, lang=self.language)
                    if page_text.strip():
                        text_parts.append(f"--- Página {i+1} ---\n{page_text}")
                except Exception as e:
                    self.logger.warning(f"Erro OCR na página {i+1}: {e}")
                    continue
            
            content = "\n\n".join(text_parts)
            return md(content) if HAS_MARKDOWN else content, True
            
        except Exception as e:
            raise FileExtractionError(f"Erro no OCR do PDF: {e}")

    def _extract_json(self, file_path: Path, ocr_func: Optional[Callable] = None) -> str:
        """Extrai e formata dados JSON"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Formata JSON de forma legível
            formatted = pprint.pformat(data, indent=2, width=100)
            
            # Converte para markdown com formatação
            markdown_content = f"# Arquivo JSON: {file_path.name}\n\n```json\n{json.dumps(data, indent=2, ensure_ascii=False)}\n```\n\n## Estrutura Formatada\n\n```\n{formatted}\n```"
            
            return markdown_content
            
        except json.JSONDecodeError as e:
            raise FileExtractionError(f"JSON malformado: {e}")
        except Exception as e:
            raise FileExtractionError(f"Erro ao processar JSON: {e}")

    # ===== MÉTODOS DE CONVERSÃO =====

    def _markdown_to_text(self, markdown_str: str) -> str:
        """Converte markdown para texto simples"""
        import re
        
        # Remove marcadores markdown
        text = re.sub(r'#{1,6}\s*', '', markdown_str)  # Headers
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)   # Bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)       # Italic
        text = re.sub(r'`(.*?)`', r'\1', text)         # Code
        text = re.sub(r'>\s*', '', text)               # Blockquotes
        text = re.sub(r'^\s*[-*+]\s*', '• ', text, flags=re.MULTILINE)  # Lists
        text = re.sub(r'^\s*\d+\.\s*', '• ', text, flags=re.MULTILINE)  # Numbered lists
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # Links
        
        # Limpa espaços extras
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()

    def _markdown_to_html(self, markdown_str: str) -> str:
        """Converte markdown para HTML"""
        if not HAS_MARKDOWN:
            self.logger.warning("Pacote 'markdown' não instalado. Retornando markdown puro.")
            return f"<pre>{markdown_str}</pre>"
        
        try:
            html = markdown.markdown(
                markdown_str,
                extensions=['tables', 'fenced_code', 'toc']
            )
            return html
        except Exception as e:
            self.logger.error(f"Erro na conversão para HTML: {e}")
            return f"<pre>{markdown_str}</pre>"

    def _default_llm_processing(self, text: str) -> str:
        """Processamento LLM padrão (stub)"""
        return text

    # ===== MÉTODOS UTILITÁRIOS =====

    def get_supported_formats(self) -> List[str]:
        """Retorna lista de formatos suportados"""
        return self.supported_formats.copy()

    def get_processing_stats(self) -> Dict[str, Any]:
        """Retorna estatísticas de processamento"""
        return {
            "supported_formats": len(self.supported_formats),
            "max_workers": self.max_workers,
            "max_file_size_mb": self.max_file_size_bytes / (1024 * 1024),
            "dependencies": {
                "pymupdf": HAS_PYMUPDF,
                "ocr": HAS_OCR,
                "docx": HAS_DOCX,
                "pandas": HAS_PANDAS,
                "markdown": HAS_MARKDOWN
            }
        }

    def validate_dependencies(self) -> Dict[str, bool]:
        """Valida todas as dependências opcionais"""
        return {
            "pymupdf": HAS_PYMUPDF,
            "ocr": HAS_OCR,
            "docx": HAS_DOCX,
            "pandas": HAS_PANDAS,
            "markdown": HAS_MARKDOWN
        }